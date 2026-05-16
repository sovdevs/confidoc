"""Markdown → DOCX export for Confidoc translation packages.

Converts the prepared (anonymized) markdown to a Word document suitable
for sending to a human translator. Pseudonymization tokens ([PATIENT_001]
etc.) are preserved as-is and highlighted for the translator's awareness.

Supports:
  # / ## / ###          → Heading 1 / 2 / 3
  - / * bullet lines    → List Bullet
  | table rows |        → Word table
  ---                   → horizontal rule (page-break between sections)
  blank line            → paragraph break
  regular text          → Normal paragraph
  YAML front matter     → skipped (between opening ---)
"""

from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_TOKEN_RE = re.compile(r'\[[A-Z][A-Z_]*_\d{3}\]')


def _add_para(doc: Document, text: str, style: str = "Normal") -> None:
    """Add a paragraph, highlighting any pseudonymization tokens."""
    para = doc.add_paragraph(style=style)
    parts = _TOKEN_RE.split(text)
    tokens = _TOKEN_RE.findall(text)
    for i, part in enumerate(parts):
        if part:
            para.add_run(part)
        if i < len(tokens):
            run = para.add_run(tokens[i])
            run.bold = True
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)   # dark red


def _add_table(doc: Document, rows: list[str]) -> None:
    """Parse pipe-separated rows into a Word table."""
    parsed: list[list[str]] = []
    for row in rows:
        cols = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cols)

    if not parsed:
        return

    n_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        # Skip separator rows (--- cells)
        if all(re.match(r'^-+$', c.strip()) for c in row if c.strip()):
            continue
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def md_to_docx(markdown: str, src_lang: str, tgt_lang: str, title: str = "") -> bytes:
    """Convert anonymized markdown to DOCX bytes."""
    doc = Document()

    # Document properties
    core = doc.core_properties
    if title:
        core.title = title
    core.keywords = "confidoc; anonymized; translation"

    # Cover note
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run(
        f"Confidoc Export — Anonymized Document\n"
        f"Source: {src_lang}  →  Target: {tgt_lang}\n"
        f"Tokens in [BRACKETS] are pseudonymization placeholders — preserve them verbatim."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()   # spacer

    lines = markdown.splitlines()
    in_front_matter = False
    in_table: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # YAML front matter
        if i == 0 and line.strip() == "---":
            in_front_matter = True
            i += 1
            continue
        if in_front_matter:
            if line.strip() == "---":
                in_front_matter = False
            i += 1
            continue

        # Flush pending table when non-table line arrives
        if in_table and not line.startswith("|"):
            _add_table(doc, in_table)
            in_table = []

        # Table row
        if line.startswith("|"):
            in_table.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            style = f"Heading {min(level, 3)}"
            doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        # Horizontal rule → page break
        if re.match(r'^---+\s*$', line):
            doc.add_page_break()
            i += 1
            continue

        # Bullet list
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            _add_para(doc, m.group(1).strip(), "List Bullet")
            i += 1
            continue

        # Numbered list
        m = re.match(r'^\d+\.\s+(.*)', line)
        if m:
            _add_para(doc, m.group(1).strip(), "List Number")
            i += 1
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        _add_para(doc, line)
        i += 1

    # Flush any remaining table
    if in_table:
        _add_table(doc, in_table)

    # Save to bytes
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
