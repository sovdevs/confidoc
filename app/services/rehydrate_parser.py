"""Zone 1 rehydration service.

Extracts embedded package_id from Confidoc-exported files and applies
token → original-PHI substitution.

Security contract
─────────────────
• All rehydration is Zone 1 only — the mapping key never leaves Zone 1.
• The caller (route) is responsible for loading and decrypting the mapping.
• This module only parses file structure and applies text substitution.

Supported formats: .md  .xliff  .sdlxliff  .tmx  .docx
Embedded marker:   confidoc-package:{package_id}
  – Text formats: XML/HTML comment or free-text in header elements
  – DOCX: stored in core_properties.keywords
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

_PKG_RE = re.compile(r'confidoc-package:([A-Za-z0-9_-]+)')


def extract_package_id(content: bytes, filename: str) -> str | None:
    """Return the embedded package_id, or None if not found."""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        return _extract_docx_package_id(content)
    # All text-based formats: MD, XLIFF, SDLXLIFF, TMX
    try:
        text = content.decode("utf-8", errors="replace")
    except Exception:
        return None
    m = _PKG_RE.search(text)
    return m.group(1) if m else None


def _extract_docx_package_id(content: bytes) -> str | None:
    try:
        from docx import Document
        doc = Document(BytesIO(content))
        keywords = doc.core_properties.keywords or ""
        m = _PKG_RE.search(keywords)
        return m.group(1) if m else None
    except Exception:
        return None


def rehydrate_content(content: bytes, filename: str, token_map: dict[str, str]) -> bytes:
    """Replace stable tokens with original PHI values.

    For DOCX: operates run-by-run to preserve formatting where possible.
    For all other formats: plain string replacement on UTF-8 text.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        return _rehydrate_docx(content, token_map)
    text = content.decode("utf-8")
    for token, original in token_map.items():
        text = text.replace(token, original)
    return text.encode("utf-8")


def _rehydrate_docx(content: bytes, token_map: dict[str, str]) -> bytes:
    from docx import Document

    doc = Document(BytesIO(content))

    def _replace_para(para):
        for run in para.runs:
            for token, original in token_map.items():
                if token in run.text:
                    run.text = run.text.replace(token, original)

    for para in doc.paragraphs:
        _replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
